from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create a new Document
doc = Document()

# Set document style
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Header
doc.add_heading('Nisarg Chauhan', 0)
doc.add_paragraph('Gandhinagar, Gujarat, India | +91 99999 99999\n'
                  'nisargchauhan777888@gmail.com | nisargchauhan777888@ieee.org\n'
                  'LinkedIn: linkedin.com/in/nisarg-chauhan-63164925b\n'
                  'GitHub: github.com/sou-nisarg')

# Professional Summary
doc.add_heading('Professional Summary', level=1)
doc.add_paragraph(
    "Passionate and driven Computer Engineering student specializing in Artificial Intelligence and Machine Learning "
    "at Silver Oak University. Experienced in developing ML/DL models, working with real-world datasets, "
    "and building impactful AI applications in domains like computer vision and natural language processing. Adept at "
    "using Python, Power BI, and other modern tools to solve analytical problems. Seeking to contribute to "
    "data-driven organizations through innovative solutions and collaborative research."
)

# Education
doc.add_heading('Education', level=1)
doc.add_paragraph(
    "Silver Oak University, Ahmedabad\n"
    "B.Tech in Computer Engineering – AI/ML Specialization\n"
    "May 2022 – June 2026\n"
    "CPI: \n"
    "Relevant Coursework: Machine Learning, Deep Learning, Computer Vision, Data Structures, Statistical Modeling"
)

# Technical Skills
doc.add_heading('Technical Skills', level=1)
doc.add_paragraph(
    "- Programming: Python, C++, Java\n"
    "- ML/DL Libraries: Scikit-learn, TensorFlow, Keras, PyTorch\n"
    "- Data Handling: Pandas, NumPy, SQL\n"
    "- Visualization Tools: Power BI, Matplotlib, Seaborn\n"
    "- Computer Vision: OpenCV, MediaPipe\n"
    "- Tools: Jupyter Notebook, Anaconda, Android Studio, VS Code\n"
    "- Cloud: Familiarity with AWS, Google Colab"
)

# Projects
doc.add_heading('Key Projects', level=1)
projects = [
    ("Car Price Prediction Interface",
     "Used regression algorithms to predict car prices based on various features.\nTechnologies: Python, Sklearn, Streamlit"),
    ("Craters & Boulders Detector for Planetary Exploration",
     "Built an image-based detector for identifying craters and boulders in space mission datasets.\nTechnologies: OpenCV, Deep Learning"),
    ("Corrosion and Fire Extinguisher Detector (Industrial AI)",
     "Created object detection models to locate corrosion and fire extinguishers in an industrial setup.\nTechnologies: YOLOv5, CV2"),
    ("Indian Hand Sign Language Interpreter",
     "Developed a real-time interpreter using hand gesture recognition.\nTechnologies: MediaPipe, CNN")
]
for title, description in projects:
    doc.add_paragraph(f"• {title}", style='List Bullet')
    doc.add_paragraph(description, style='Normal')

# Internship Experience
doc.add_heading('Internship Experience', level=1)
doc.add_paragraph(
    "Sahana System Limited\n"
    "AI Intern — Oct 2024 – Dec 2024\n"
    "Data Management Intern — Jun 2024 – Jul 2024\n"
    "- Engaged in real-time ML projects and data pipelines\n"
    "- Built data visualization dashboards and trained ML models\n\n"
    "Batwebs.com\n"
    "Data Science Intern — Dec 2023 – Jan 2024\n"
    "- Worked on classification, regression, and clustering\n"
    "- Created reports using Power BI\n"
    "- Explored supervised and unsupervised ML concepts\n"
    "- Project GitHub: https://github.com/sou-nisarg/Batwebs_nisarg_repo.git"
)

# Leadership & Extracurricular Activities
doc.add_heading('Leadership & Extracurricular Activities', level=1)
doc.add_paragraph(
    "IEEE Silver Oak University Student Branch\n"
    "- Chairperson, IEEE SPS SBC — Jan 2025 – Present\n"
    "- Global IEEE SPS Ambassador — Feb 2025 – Present\n"
    "- Secretary, IEEE SOU SIGHT — Jan 2024 – Dec 2024\n"
    "- Organized multiple workshops, technical sessions, and outreach activities"
)

# Certifications (Placeholder)
doc.add_heading('Certifications', level=1)
doc.add_paragraph(
    "- Deep Learning Specialization – Coursera\n"
    "- AI For Everyone – Andrew Ng"
)

# Languages
doc.add_heading('Languages', level=1)
doc.add_paragraph(
    "- English – Professional\n"
    "- Hindi – Fluent\n"
    "- Gujarati – Native"
)

# Interests
doc.add_heading('Interests', level=1)
doc.add_paragraph(
    "- Computer Vision & NLP\n"
    "- LLMs and Generative AI\n"
    "- AI for Social Good\n"
    "- E-Sports and Gaming Tech"
)

# Save the document
file_path = "/mnt/data/Nisarg_Chauhan_Resume_2025.docx"
doc.save(file_path)

file_path  # Return path for download or further processing
